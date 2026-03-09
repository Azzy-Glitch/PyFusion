from docx import Document

doc = Document()
doc.add_heading('Is-A and Has-A Relationship in C++', level=1)

doc.add_heading('1. Is-A Relationship', level=2)
doc.add_paragraph(
"Definition:\n"
"An Is-A relationship means one class is a type of another class. "
"It is implemented using inheritance in C++.\n\n"
"In simple words: If we can say 'A is a type of B', then it is an Is-A relationship.\n\n"
"Examples:\n"
"- A Dog is an Animal\n"
"- A Car is a Vehicle\n"
"- A Student is a Person"
)

doc.add_heading('Example in C++', level=3)
doc.add_paragraph(
'#include <iostream>\n'
'using namespace std;\n\n'
'class Animal\n'
'{\n'
'public:\n'
'    void eat()\n'
'    {\n'
'        cout << "Animal can eat\\n";\n'
'    }\n'
'};\n\n'
'class Dog : public Animal\n'
'{\n'
'public:\n'
'    void bark()\n'
'    {\n'
'        cout << "Dog can bark\\n";\n'
'    }\n'
'};\n\n'
'int main()\n'
'{\n'
'    Dog d;\n'
'    d.eat();\n'
'    d.bark();\n'
'}'
)

doc.add_heading('Explanation', level=3)
doc.add_paragraph(
"Dog inherits from Animal, meaning Dog IS-A Animal. "
"So the Dog class can use functions of Animal like eat()."
)

doc.add_heading('2. Has-A Relationship', level=2)
doc.add_paragraph(
"Definition:\n"
"A Has-A relationship means a class contains another class as a member.\n\n"
"In simple words: If we can say 'A has a B', then it is a Has-A relationship.\n\n"
"Examples:\n"
"- A Car has an Engine\n"
"- A Computer has a CPU\n"
"- A Person has a Heart"
)

doc.add_heading('Example in C++', level=3)
doc.add_paragraph(
'#include <iostream>\n'
'using namespace std;\n\n'
'class Engine\n'
'{\n'
'public:\n'
'    void start()\n'
'    {\n'
'        cout << "Engine started\\n";\n'
'    }\n'
'};\n\n'
'class Car\n'
'{\n'
'public:\n'
'    Engine e;\n\n'
'    void drive()\n'
'    {\n'
'        e.start();\n'
'        cout << "Car is moving\\n";\n'
'    }\n'
'};\n\n'
'int main()\n'
'{\n'
'    Car c;\n'
'    c.drive();\n'
'}'
)

doc.add_heading('Explanation', level=3)
doc.add_paragraph(
"The Car class contains an Engine object. "
"This means Car HAS-A Engine."
)

doc.add_heading('Difference Between Is-A and Has-A', level=2)
doc.add_paragraph(
"Is-A Relationship:\n"
"- Implemented using inheritance\n"
"- Represents 'is a type of'\n"
"- Example: Dog is an Animal\n\n"
"Has-A Relationship:\n"
"- Implemented by creating objects of another class\n"
"- Represents 'contains'\n"
"- Example: Car has an Engine"
)

# Save the document
file_path = "is_a_has_a_relationship_cpp.docx"
doc.save(file_path)

print("Document created successfully.")
print("File saved as:", file_path)